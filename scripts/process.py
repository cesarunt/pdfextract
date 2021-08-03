import os, re
import datetime
import spacy
# import docx
import pandas as pd
from bs4 import BeautifulSoup
from utils.process_pdf import *
from utils.process_data import *
from utils.config import cfg
from docx import Document

def build_document(title, text_pdf):
    document = Document() 

    # add a header
    document.add_heading(title)

    # add a paragraphs
    for item in text_pdf:
        print("type: "+str(type(item)))
        document.add_paragraph(str(item))

    # add a paragraph within an italic text then go on with a break.
    # paragraph = document.add_paragraph()
    # run = paragraph.add_run()
    # run.italic = True
    # run.add_text("text will have italic style")
    # run.add_break()
    
    return document

def clear_report():
    # open("output/report.html", "w").close()
    open(cfg.FILES.SINGLE_OUTPUT+"/background.txt", "w").close()

def addText_background(line):
    if line != "":
        text_pdf.append(line + "\n")

def addText_background_(line, limit=""):
    if limit == "":
        text_pdf.append(line + " ")
    else:
        text_pdf.append(line + limit)

def pdf_process():
    clear_report()
    fname = os.listdir(cfg.FILES.SINGLE_SPLIT+"/")
    fname.sort(key=lambda f: int(re.sub('\D', '', f)))
    length = len(fname)

    global text_pdf
    text_pdf = []

    page = 0
    language = ""
    text_page = ""
    band_method = False
    band_title = False
    title_text = ""         # 01. FIND THE TITLE TEXT
    authors_name = []       # 02. FIND THE AUTHORS NAME
    year = 0                # 03. FIND THE PUBLISHING YEAR
    objective = ""          # 04. FIND THE PATTERN OBJECTIVE
    text_methodology = []   # 05. FIND THE METHODOLOGY 
    type_level = ""         #  -. FIND THE PATTERN TYPE 
    design = ""             #  -. FIND THE PATTERN DESIGN
    approach = ""           #  -. FIND THE PATTERN APPROACH
    samples = ""            # 06. FIND THE PATTERN SAMPLES
    tools = ""              # 07. FIND THE PATTERN TOOLS
    text_results = []       # 08. FIND THE PATTERN RESULTS
    text_conclusions = []   # 09. FIND THE PATTERN CONCLUSIONS

    article_res = False
    band_result = False
    result_res = False
    listQuan = []
    listQual = []
    methods0     = ""
    results0     = ""
    conclusions0 = ""

    result_process = False
    PATTERN_OBJE = []
    level_appl = False; level_pred = False; level_expi = False; level_rela = False; level_desc = False; level_expo = False

    for page in range(length): #Repeat each operation for each document.
        print("Page 0"+str(int(page+1)) +": "+fname[page])

        # 1. FIND THE TITLE TEXT
        # ============================================================================================
        # - Extract text with functions PDFminer
        text_page = convert_pdf_to_text((cfg.FILES.SINGLE_SPLIT+'/{}').format(fname[page])) #Extract text with PDF_to_text Function call
        text_html = convert_pdf_to_html((cfg.FILES.SINGLE_SPLIT+'/{}').format(fname[page])) #Extract text with PDF_to_html Function call
        # text_html_out = text_html.decode("utf-8")     #Decode result from bytes to text
        # print(text_html)

        if text_page != "" and band_title == False:
            # - Using BeautifulSoup to parse the text
            soup = BeautifulSoup(text_html, 'html.parser')
            patt = re.compile("font-size:(\d+)")
            text_parser = [(tag.text.strip(), int(patt.search(tag["style"]).group(1))) for tag in soup.select("[style*=font-size]")] 
            # - Getting the max value of the font
            title_font  = max(text_parser, key=lambda x:x[1])[1]
            title_list_text = []
            if (title_font==55) : title_font=15
            title_list_text = [key for key, value in text_parser if value == title_font]
            title_text = (''.join(title_list_text)).replace("\n", "")
            if title_font > 10 and title_font < 30 and len(title_text) > 1:
                title_text = title_text.split("___")[0]
                # if "Resumen" in title_text :
                    # title_text =(' '.join(title_list_text)).split("Resumen")[0]
                band_title = True
            # - RESULT: "title_text"
            # ============================================================================================
            # 2. FIND THE AUTHOR NAME
            # ============================================================================================
            if band_title == True or page == 0:
                # - Getting the language of text
                if language == "" :
                    # get the text's language and patterns for NLP
                    language = lang_getLanguage(text_page)
                    lib_spacy, STOP_WORDS, patterns, patterns_level, patterns_approach = lang_loadPatterns(language)
                    BLOCK_WORDS, PATTERN_OBJE, PATTERN_METH, PATTERN_TYPE, PATTERN_DESI, PATTERN_APPR, PATTERN_LEVE, PATTERN_SAMP, PATTERN_TOOL, PATTERN_RESU, PATTERN_CONC = patterns
                    PATTERN_LEVE_APPL, PATTERN_LEVE_PRED, PATTERN_LEVE_EXPI, PATTERN_LEVE_RELA, PATTERN_LEVE_DESC, PATTERN_LEVE_EXPO = patterns_level
                    PATTERN_APPR_QUAN, PATTERN_APPR_QUAL = patterns_approach      
                NLP = spacy.load(lib_spacy)
                text_nlp = NLP(text_page)
                person_label = []
                person_tot = 0
                title_inc = 0
                title_tot = 0
                w = 0
                
                if len(authors_name) == 0 :
                    print("List Original NLP items : " + str(len(text_nlp.ents)))
                    for word in text_nlp.ents:
                        word_ = word.text.lower().split(" ")[0]
                        if  word_ not in STOP_WORDS:
                            word_formated  = format_word_dash(word.text)
                            size_word = len(word_formated.split())
                            # print(str(w)+" .-. "+str(size_word)+" .-. "+word_formated+"- "+word.label_)
                            if len(word_formated) > 0 :
                                title_inc = find_word_in_title(word_formated, title_text)
                                # print(str(title_tot)+" .-. "+str(title_inc)+" .-. "+word_formated+"- "+word.label_) ##############
                                result = find_words_allowed(word_formated)
                                if result == True :
                                    word_formated = clear_word(word_formated, BLOCK_WORDS)
                                    authors_name.append(tuple([w, 0, word.label_, word_formated]))
                                else :
                                    if title_inc == 1 :
                                        title_tot += 1
                                    if title_tot > 0 and title_inc==0 :
                                        # if (word.label_ != "PER" and person_tot>0) or (word.label_ == "PER" and word_ in BLOCK_WORDS) :
                                        if word_ in BLOCK_WORDS:
                                            break
                                        if (word.label_ == "PER" and size_word > 1) :
                                            person_tot += 1
                                            word_formated = clear_word(word_formated, BLOCK_WORDS)
                                            authors_name.append(tuple([w, title_inc, word.label_, word_formated]))
                                    if title_tot == 0 and title_inc == 0:
                                        if (word.label_ != "PER" and person_tot>0) or (word.label_ == "PER" and word_ in BLOCK_WORDS):
                                            # print("continue")
                                            w += 1
                                            continue
                                        if (word.label_ == "PER" and size_word > 1 and len(word_)>2 and word_[0:4]!="http") :
                                            person_tot += 1
                                            word_formated = clear_word(word_formated, BLOCK_WORDS)
                                            authors_name.append(tuple([w, title_inc, word.label_, word_formated]))
                            w += 1
             # - RESULT: "authors_name"

                if article_res == False:        article_res,_   = getData_ResultText(text_page, ['artículo'])
             # ============================================================================================
             # 3. FIND THE DOCUMENT PUBLISHING YEAR 
             # ============================================================================================
             # if text_page != "" :
                # - Getting the max year of PDF 1st page
        
        if year == 0 or year < 1000:
            list_year = re.findall("(\d{4})",text_page)
            year_max = 0
            date = datetime.date.today()
            for year in list_year :
                year = int(year)
                if year>year_max and year<=date.year :
                    year_max = year
            year = year_max
             # - RESULT: "year"
             # ============================================================================================    
                 # print("TIT: "+title_text)
        
        if objective == "":     objective = getData_LongText(text_page, PATTERN_OBJE, 'E', '. '); #print(objective)
        
        text_method = ""
        method_pos = 0
        if text_page != "" and language!="" and band_method == False:
            method_res, method_pos        = getData_ResultText(text_page, PATTERN_METH)
            if method_res :
                band_method = True
                method_page = page + 1
        
        if text_page != "" and language!="" and band_method == True :
            result_res, result_pos       = getData_ResultText(text_page[method_pos:], PATTERN_RESU)
            if result_res == False :
                # text_method = text_page[method_pos:-1].replace("\n", " ")
                text_method = text_page[method_pos:-1]
                # print("METHOD")
                # print(text_method)
            if result_res == True:
                result_page = page + 1
                band_result = True
                if method_page < result_page :
                    # text_method = text_page[0:result_pos].replace("\n", "")
                    text_method = text_page[0:result_pos]
                else :
                    # text_method = text_page[method_pos:result_pos].replace("\n", " ")
                    text_method = text_page[method_pos:result_pos]
                text_result = text_page[result_pos:]
            
            if text_method != "" :
                text_con = getData_LongText(text_method, PATTERN_METH, 'E', '.')
                if text_con != "" :
                    text_methodology.append(text_con)
                
            # - Getting the methodology (methodology, design, approach, level)
            if type_level == "" :   type_level  = getData_LongText(text_method, PATTERN_TYPE, 'S', ', ')
            if design == "" :       design     = getData_LongText(text_method, PATTERN_DESI, 'S', ', ')
            if approach == "" :     approach  = getData_LongText(text_method, PATTERN_APPR, 'S', '. ')
            # - getting 2 approach
            # if approach_quan == False : approach_quan = getTools_ResultCount(text_method, PATTERN_APPR_QUAN)
            listQn = getTools_ResultCount(text_method, PATTERN_APPR_QUAN)
            if len(listQn) > 0 :  listQuan = listQuan + listQn
            listQl = getTools_ResultCount(text_method, PATTERN_APPR_QUAL)
            if len(listQl) > 0 :  listQual = listQual + listQl
            # if approach_qual == False : listQual = getTools_ResultCount(text_method, PATTERN_APPR_QUAL)
            # - getting 5 levels
            if level_appl == False :    level_appl = getLevel_Result(text_method, PATTERN_LEVE_APPL)
            if level_pred == False :    level_pred = getLevel_Result(text_method, PATTERN_LEVE_PRED)
            if level_expi == False :    level_expi = getLevel_Result(text_method, PATTERN_LEVE_EXPI)
            if level_rela == False :    level_rela = getLevel_Result(text_method, PATTERN_LEVE_RELA)
            if level_desc == False :    level_desc = getLevel_Result(text_method, PATTERN_LEVE_DESC)
            if level_expo == False :    level_expo = getLevel_Result(text_method, PATTERN_LEVE_EXPO)
            
            if text_page != "" and language!="" and tools == "":
                # - Getting the tools
                tools = getData_Long(text_method, PATTERN_TOOL)
            
            if text_page != "" and language!="" and samples == "" and band_method==True:
                # samples   = getData_Long(text_page, PATTERN_SAMP)
                samples   = getData_LongText(text_method, PATTERN_SAMP, 'E', '.\n')

            ## LOS RESULTADOS SE DEBEN OBTENER DE LA MISMA FORMA QUE LA METODOLOGIA
            if text_page != "" and language!="" and band_result == True :
                result_res, result_pos = getData_ResultText(text_result, PATTERN_RESU)
                # - Getting the text_results
                if result_res :
                    text_con = getData_LongText(text_result, PATTERN_RESU, 'E', '.\n')
                    if text_con != "" :
                        text_results.append(text_con.replace("\n", "")) 

                    # text_res = getData_LongText_Result(text_result, PATTERN_RESU, 'E', '.\n')
                    # for item in text_res :
                    #     text_results.append(item.replace("\n", ""))

                    band_result = False
            
            # if text_page != "" and language!="" :
            #     text_con = getData_LongText(text_page, PATTERN_CONC, 'E', '. ')
            #     if text_con != "":
            #         text_conclusions.append(text_con)
                # text_conclusions = text_conclusions.replace("\n", "")
                # text_conclusions    = getData_LongText(text_page, PATTERN_CONC, 'E', '.')
                # print("conclusiones: "+text_conclusions[page])

        if text_page != "" and language!="" and page<=1 :
            text_con = getData_LongText(text_page, PATTERN_METH, 'E', '.')
            if text_con != "" :
                text_methodology.append(text_con)

        if text_page != "" and language!="" and samples == "":
            if band_method==True :
                samples   = getData_LongText(text_method, PATTERN_SAMP, 'E', '.\n')
            else:
                samples   = getData_LongText(text_page, PATTERN_SAMP, 'E', '.\n')
        if text_page != "" and language!="" and page<=1 :
            text_con = getData_LongText(text_page, PATTERN_RESU, 'E', '.\n')
            if text_con != "" :
                text_results.append(text_con.replace("\n", "")) 
        if text_page != "" and language!="" :
            text_con = getData_LongText(text_page, PATTERN_CONC, 'E', '.\n')
            if text_con != "" :
                text_conclusions.append(text_con)

        # ============================================================================================
        # text_pdf = []
        if band_title == True and len(authors_name)>0 and year!="" and page == length-1 :

            addText_background(fname[page].split(".")[0]+".")
            # print("------------------- END PROCESS -------------------")

            print("\n01._ AUTHOR NAME (" + str(len(authors_name))+ "):")
            for item in authors_name:
                addText_background_(item[3], ", ")
                print(item)

            print("\n02._ PUBLISHING YEAR: " + str(year))
            addText_background_('('+str(year)+') en su estudio titulado "')

            addText_background_(title_text+'" ')
            print("\n03._ TITLE TEXT (" + str(title_font)+"px):\n" + title_text)

            if article_res : text_article = "(Artículo científico)."
            else : text_article = "(Revista científica)."
            addText_background(text_article+'\n')

            addText_background("El objetivo ... " + objective)
            print("\n04._ OBJECTIVE :\n" + objective)

            addText_background("\nLa metodología ...")
            addText_background(str(text_methodology))
            print("\n05._ METHODOLOGY :\n" + str(text_methodology), end="")
            method_list = {'type_level':'tipo', 'design':'diseño', 'approach':'enfoque'}
            for key, value in method_list.items() : 
                if value in str(text_methodology) or value.capitalize() in str(text_methodology):
                    print("\n   .-"+key.capitalize()+" : " + vars()[key])
                # vars()[item]
            addText_background_("\n - Enfoque detallado: ")
            print("\n  5.1._ APPROACH DETAILS : ", end="")
            listQuan = list(dict.fromkeys(listQuan))
            if len(listQuan) > 0 :
                addText_background_("Cuantitativo, ")
                print("Cuantitativo", end=", ")
            listQual = list(dict.fromkeys(listQual))
            if listQual : 
                addText_background_("Cualitativo")
                print("Cualitativo", end="")
            addText_background_("\n - Nivel detallado: ")
            print("\n\n  5.2._ LEVEL DETAILS: ", end="")
            if level_appl : addText_background_("Aplicado, ");     print("Aplicado", end=", ")
            if level_pred : addText_background_("Predictivo, ");   print("Predictivo", end=", ")
            if level_expi : addText_background_("Explicativo, ");  print("Explicativo", end=", ")
            if level_rela : addText_background_("Relacional, ");   print("Relacional", end=", ")
            if level_desc : addText_background_("Descriptivo, ");  print("Descriptivo", end=", ")
            if level_expo : addText_background_("Exploratorio"); print("Exploratorio", end="")

            addText_background("\n\nLa muestra ... ")
            addText_background(samples)
            print("\n\n06._ SAMPLES :\n" + samples)
            
            addText_background("Técnica(s) de recolección de datos empleada(s):")
            print("\n07._ TOOLS : Técnica(s) de recolección de datos empleada(s): ")
            if len(listQuan) > 0 : addText_background("   "+str(listQuan)); print("  " + str(listQuan))
            if len(listQual) > 0 : addText_background("   "+str(listQual)); print("  " + str(listQual))

            addText_background("\nLos resultados ... ")
            print("\n08._ RESULTS : \n" + str(text_results))
            for item in text_results:
                addText_background(str("- "+item))

            addText_background("\nLas conclusiones ... ")
            # addText_background(str(text_conclusions))
            for item in text_conclusions:
                addText_background(str("- "+item))
            print("\n09._ CONCLUSIONS :\n" + str(text_conclusions))

            addText_background('\n\n')
            addText_background('REFERENCIAS')
            for item in authors_name:
                addText_background_(item[3], ", ")
            # addText_background(str(authors_name))
            addText_background_("("+str(year)+"). ")
            addText_background_(title_text.capitalize())
            addText_background_(". Obtenido de ... https://")
            # break

            document = build_document(fname[page], text_pdf)
            document.save(cfg.FILES.SINGLE_OUTPUT+'/background.docx')

            #Save extracted text to file.txt
            text_pdf_out = text_pdf #.decode("utf-8")
            with open(cfg.FILES.SINGLE_OUTPUT+"/background.txt", "a", encoding="utf-8") as text_file:
                # text_file.writelines("Page_0" + str(int(page+1)) + "\n")
                text_file.writelines(text_pdf_out)
                result_process = True
        # input("................... press enter ........................")
    
    return result_process