import re
import unicodedata
from docx import Document

ILLEGAL_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]")
# document = Document()

def castLetter(text):
    text = re.sub(r'a&#769;', 'á', text)
    text = re.sub(r'A&#769;', 'Á', text)
    text = re.sub(r'e&#769;', 'é', text)
    text = re.sub(r'E&#769;', 'É', text)
    text = re.sub(r'i&#769;', 'í', text)
    text = re.sub(r'I&#769;', 'Í', text)
    text = re.sub(r'o&#769;', 'ó', text)
    text = re.sub(r'O&#769;', 'Ó', text)
    text = re.sub(r'u&#769;', 'ú', text)
    text = re.sub(r'U&#769;', 'Ú', text)
    text = re.sub(r'&#8220;', '"', text)
    text = re.sub(r'&#8221;', '"', text)
    text = re.sub(r'n&#771;', 'ñ', text)
    text = re.sub(r'N&#771;', 'Ñ', text)
    return text

def castPDFsMT(text_schemes):
    det_title = ""
    det_author = ""
    det_year = ""
    det_magazine = ""
    det_volumen = ""
    det_link = ""
    det_page = ""
    text_scheme = []
    pdfs = dict()

    for detail in text_schemes :
        pdfs[detail['det_name']] = detail['det_value']
        if detail['det_name'] == 'enlace':
            size_doc = detail['det_npage']

    if (pdfs['título']):
        det_title     = pdfs['título'].replace('\n', ' ').replace('\r', '')
    if (pdfs['autor']):
        det_author    = pdfs['autor'].replace('\n', ' ').replace('\r', '')
    if (pdfs['año']):
        det_year      = pdfs['año'].replace('\n', ' ').replace('\r', '')
    if ('revista' in pdfs):
        det_magazine = pdfs['revista'].replace('\n', ' ').replace('\r', '')
    if ('volumen' in pdfs):
        det_volumen  = pdfs['volumen'].replace('\n', ' ').replace('\r', '')
    if ('enlace' in pdfs):
        det_link     = pdfs['enlace'].replace('\n', ' ').replace('\r', '')
    if ('página' in pdfs):
        det_page     = pdfs['página'].replace('\n', ' ').replace('\r', '')
    
    if (size_doc>40):
        # TESIS
        text_scheme.append(tuple(["N", str(det_author) + " ("+str(det_year)+"). " ]))
        text_scheme.append(tuple(["K", '"' + str(det_title) + '"']))
        if len(det_link)>0:
            text_scheme.append(tuple(["N", '.\n ' + str(det_link) ]))
    else:
        # REVISTAS
        text_scheme.append(tuple(["N", str(det_author) + " ("+str(det_year)+"). " ]))
        text_scheme.append(tuple(["K", '"' + str(det_title) + '"']))
        if (len(det_magazine)>0):
            text_scheme.append(tuple(["K", ', ' + str(det_magazine) ]))
        if (len(det_volumen)>0):
            text_scheme.append(tuple(["K", ', ' + str(det_volumen) ]))
        if (len(det_page)>0):
            text_scheme.append(tuple(["N", ', ' + str(det_page) ]))
        if len(det_link)>0:
            text_scheme.append(tuple(["N", '.\n ' + str(det_link) ]))
    
    return text_scheme
    
def readTextSchemesA(text_schemes, document, i):
    title_band = False
    text_scheme = []
    pdfs = dict()

    for detail in text_schemes :
        pdfs[detail['det_name']] = detail['det_value']
        if (str(pdfs['título'])!='' and title_band==False):
            document.add_heading(str(int(i+1)) + ". " + str(pdfs['título'].replace('\n', ' ').replace('\r', '')).capitalize())
            title_band = True
        
    # FOR SCHEME
    if (pdfs['autor'] and pdfs['año']):
        det_author    = pdfs['autor'].replace('\n', ' ').replace('\r', '')
        det_year      = pdfs['año'].replace('\n', ' ').replace('\r', '')
        text_scheme.append(tuple(["N", str(det_author) + " ("+str(det_year)+")" ]))
    if (pdfs['título']):
        det_title     = pdfs['título'].replace('\n', ' ').replace('\r', '')
        text_scheme.append(tuple(["N", ". en su investigación titulada "]))
        text_scheme.append(tuple(["K", '"' + str(det_title) + '"']))
    if ('objetivo' in pdfs):
        det_objective   = pdfs['objetivo'].replace('\n', ' ').replace('\r', '')
        if det_objective:
            text_scheme.append(tuple(["N", ". El objetivo de estudio fue"]))
            text_scheme.append(tuple(["N", " " + str(det_objective)]))
    if ('enfoque' in pdfs):
        det_approach  = pdfs['enfoque'].replace('\n', ' ').replace('\r', '')
        if det_approach:
            text_scheme.append(tuple(["N", ". A nivel metodológico la investigación fue de enfoque"]))
            text_scheme.append(tuple(["N", " " + str(det_approach)]))
    if ('diseño' in pdfs):
        det_design    = pdfs['diseño'].replace('\n', ' ').replace('\r', '')
        if det_design:
            text_scheme.append(tuple(["N", ", con un diseño"]))
            text_scheme.append(tuple(["N", " " + str(det_design)]))
    if ('nivel' in pdfs):
        det_level     = pdfs['nivel'].replace('\n', ' ').replace('\r', '')
        if det_level:
            text_scheme.append(tuple(["N", ", de nivel"]))
            text_scheme.append(tuple(["N", " " + str(det_level)]))
    if ('muestra' in pdfs):
        det_sample    = pdfs['muestra'].replace('\n', ' ').replace('\r', '')
        if det_sample:
            text_scheme.append(tuple(["N", ", la muestra se conformó por"]))
            text_scheme.append(tuple(["N", " " + str(det_sample)]))
    if ('instrumentos' in pdfs):
        det_tools     = pdfs['instrumentos'].replace('\n', ' ').replace('\r', '')
        if det_tools:
            text_scheme.append(tuple(["N", ", y se aplicaron como instrumentos"]))
            text_scheme.append(tuple(["N", " " + str(det_tools)]))
    if ('resultados' in pdfs):
        det_results    = pdfs['resultados'].replace('\n', ' ').replace('\r', '')
        if det_results:
            text_scheme.append(tuple(["N", ". Los principales resultados fueron"]))
            text_scheme.append(tuple(["N", " " + str(det_results)]))
    if ('conclusiones' in pdfs):
        det_conclussions = pdfs['conclusiones'].replace('\n', ' ').replace('\r', '')
        if det_conclussions:
            text_scheme.append(tuple(["N", ". Se concluye que"]))
            text_scheme.append(tuple(["N", " " + str(det_conclussions)]))

    return text_scheme, document

def readTextSchemesMT(text_schemes, document, i):
    title_band = False
    text_scheme = []
    pdfs = dict()

    for detail in text_schemes :
        pdfs[detail['det_name']] = detail['det_value']
        if (str(pdfs['título'])!='' and title_band==False):
            document.add_heading(str(int(i+1)) + ". " + str(pdfs['título'].replace('\n', ' ').replace('\r', '')).capitalize())
            title_band = True
        if len(str(detail['det_name']).split()) > 1 and detail['det_value']!='':
            # Subtitle for Atrrtibute -> AA
            text_scheme.append(tuple(["AS", '' + str(detail['det_name']).title() + '\n']))
            # Text for Atrrtibute -> AT
            text_scheme.append(tuple(["AT", '' + str(detail['det_value']) + '\n']))

    return text_scheme, document

def build_pdfA(text_schemes, document = None, i = 1):
    if document == None:
        document = Document()
    else:
        i = i + 1

    text_scheme, document = readTextSchemesA(text_schemes['foundlist'], document, int(i - 1))
    p = document.add_paragraph()

    for key, value in text_scheme:
        text = value.replace('\n', ' ').replace('\r', '')
        try:
            text = unicode(text, 'utf-8')
        except NameError:
            pass
        text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
        text = castLetter(text)
        html = ILLEGAL_XML_CHARS_RE.sub("", text)
        line = p.add_run(str(html))
        if key == "K": line.italic = True; p.add_run("\n")

    p.add_run("\n")
    p.add_run("\n")
    line = p.add_run("Referencias")
    line.bold = True

    # For A
    text_scheme = []
    text_scheme = castPDFsMT(text_schemes['foundlist'])
    p = document.add_paragraph()

    for key, value in text_scheme:
        text = value.replace('\n', ' ').replace('\r', '')
        try:
            text = unicode(text, 'utf-8')
        except NameError:
            pass
        text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
        text = castLetter(text)
        html = ILLEGAL_XML_CHARS_RE.sub("", text)
        line = p.add_run(str(html))
        if key == "K": line.italic = True
        # if key == "L": line.color = "#CB7825"; p.add_run("\n")
            
    return document, i

def build_pdfMT(text_schemes, document = None, i = 1):
    if document == None:
        document = Document()
    else:
        i = i + 1
    
    text_scheme, document = readTextSchemesMT(text_schemes['foundlist'], document, int(i - 1))
    p = document.add_paragraph()

    for key, value in text_scheme:
        text = value.replace('\n', ' ').replace('\r', '')
        try:
            text = unicode(text, 'utf-8')
        except NameError:
            pass
        text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
        text = castLetter(text)
        html = ILLEGAL_XML_CHARS_RE.sub("", text)
        line = p.add_run(str(html))
        if key == "AS": line.bold = True; p.add_run("\n")
        elif key == "AT": line.bold = False; p.add_run("\n"); p.add_run("\n")
        elif key == "K": line.italic = True; p.add_run("\n")
    
    p.add_run("\n")
    line = p.add_run("Referencias")
    line.bold = True

    # For MT
    text_scheme = []
    text_scheme = castPDFsMT(text_schemes['foundlist'])
    p = document.add_paragraph()

    for key, value in text_scheme:
        text = value.replace('\n', ' ').replace('\r', '')
        try:
            text = unicode(text, 'utf-8')
        except NameError:
            pass
        text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
        text = castLetter(text)
        html = ILLEGAL_XML_CHARS_RE.sub("", text)
        line = p.add_run(str(html))
        if key == "K": line.italic = True
        if key == "L": line.color = "#CB7825"; p.add_run("\n")
            
    return document, i

def build_project(text_schemes, pro_type):
    document = Document() 
    text_scheme = []

    i = 0
    for key, details in text_schemes.items():
        text_scheme = []
        # Get Data from "Antecedente"
        if details['type'] == "A" and pro_type in ("T", "A"):
            text_scheme, document = readTextSchemesA(details['details'], document, i)
            i = i + 1
            p = document.add_paragraph()

            for key, value in text_scheme:
                text = value.replace('\n', ' ').replace('\r', '')
                try:
                    text = unicode(text, 'utf-8')
                except NameError:
                    pass
                text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
                text = castLetter(text)
                html = ILLEGAL_XML_CHARS_RE.sub("", text)
                line = p.add_run(str(html))
                if key == "K": line.italic = True; p.add_run("\n")

            p.add_run("\n")
            p.add_run("\n")
            line = p.add_run("Referencias")
            line.bold = True

            # For A
            text_scheme = []
            text_scheme = castPDFsMT(details['details'])
            p = document.add_paragraph()

            for key, value in text_scheme:
                text = value.replace('\n', ' ').replace('\r', '')
                try:
                    text = unicode(text, 'utf-8')
                except NameError:
                    pass
                text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
                text = castLetter(text)
                html = ILLEGAL_XML_CHARS_RE.sub("", text)
                line = p.add_run(str(html))
                if key == "K": line.italic = True

        if details['type'] == "M" and pro_type in ("T", "M"):
            text_scheme, document = readTextSchemesMT(details['details'], document, i)
            i = i + 1
            p = document.add_paragraph()

            for key, value in text_scheme:
                text = value.replace('\n', ' ').replace('\r', '')
                try:
                    text = unicode(text, 'utf-8')
                except NameError:
                    pass
                text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
                text = castLetter(text)
                html = ILLEGAL_XML_CHARS_RE.sub("", text)
                line = p.add_run(str(html))
                if key == "AS": line.bold = True; p.add_run("\n")
                elif key == "AT": line.bold = False; p.add_run("\n"); p.add_run("\n")
                elif key == "K": line.italic = True; p.add_run("\n")
            
            p.add_run("\n")
            line = p.add_run("Referencias")
            line.bold = True

            # For MT
            text_scheme = []
            text_scheme = castPDFsMT(details['details'])
            p = document.add_paragraph()

            for key, value in text_scheme:
                text = value.replace('\n', ' ').replace('\r', '')
                try:
                    text = unicode(text, 'utf-8')
                except NameError:
                    pass
                text = unicodedata.normalize('NFD', text).encode('ascii', 'xmlcharrefreplace').decode("utf-8")
                text = castLetter(text)
                html = ILLEGAL_XML_CHARS_RE.sub("", text)
                line = p.add_run(str(html))
                if key == "K": line.italic = True
                if key == "L": line.color = "#CB7825"; p.add_run("\n")

    return document